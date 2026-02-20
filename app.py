import io
import os
import re
import tempfile
import subprocess
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, date
from zoneinfo import ZoneInfo

import psycopg2
from psycopg2.extras import RealDictCursor
from flask import Flask, render_template, request, redirect, url_for, send_file
from werkzeug.utils import secure_filename

from docx import Document
from docx.shared import Mm
from PIL import Image

# -----------------------------
# Config
# -----------------------------
APP_TZ = ZoneInfo(os.environ.get("APP_TZ", "America/Sao_Paulo"))
RETENCAO_DIAS = int(os.environ.get("RETENCAO_DIAS", "10"))

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROPOSTA_TEMPLATE = os.path.join(BASE_DIR, "template.docx")
CONTRATO_TEMPLATE = os.path.join(BASE_DIR, "contrato_template.docx")

# -----------------------------
# Flask
# -----------------------------
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-key-change-me")


# -----------------------------
# Helpers: formatação e parsing
# -----------------------------
MESES = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
]

UNIDADES = ["zero", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
DEZ_A_DEZENOVE = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
DEZENAS = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
CENTENAS = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]


def capitalizar_primeira(s: str) -> str:
    s = (s or "").strip()
    return s[:1].upper() + s[1:] if s else ""


def formatar_inteiro_ptbr(n: int) -> str:
    # 1000 -> 1.000
    return f"{n:,}".replace(",", ".")


def formatar_decimal_ptbr(d: Decimal) -> str:
    d = d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    s = f"{d:,.2f}"  # 1,234.56
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def parse_valor_decimal(s: str) -> Decimal:
    if s is None:
        raise ValueError("Valor inválido")
    s = s.strip()
    if not s:
        raise ValueError("Valor inválido")
    s = s.replace("R$", "").replace("r$", "").replace(" ", "")
    # PT-BR: vírgula é decimal; ponto é milhar
    if "," in s:
        s = s.replace(".", "").replace(",", ".")
    if not re.fullmatch(r"-?\d+(\.\d+)?", s):
        raise ValueError("Valor inválido")
    return Decimal(s).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def parse_int(s: str) -> int:
    digits = re.sub(r"\D", "", s or "")
    if not digits:
        raise ValueError("Número inválido")
    return int(digits)


def parse_data_digitos(s: str) -> date:
    """
    Aceita:
      - 20022026 (DDMMAAAA)
      - 20/02/2026
      - 200226 (DDMMAA) -> assume 20AA
      - 20/02/26
    """
    digits = re.sub(r"\D", "", s or "")
    if len(digits) == 8:
        dd = int(digits[0:2]); mm = int(digits[2:4]); yyyy = int(digits[4:8])
    elif len(digits) == 6:
        dd = int(digits[0:2]); mm = int(digits[2:4]); yy = int(digits[4:6])
        yyyy = 2000 + yy
    else:
        raise ValueError("Data inválida (use DDMMAAAA)")
    return date(yyyy, mm, dd)


def data_por_extenso(d: date, mes_capitalizado: bool = False) -> str:
    mes = MESES[d.month - 1]
    if mes_capitalizado:
        mes = mes[:1].upper() + mes[1:]
    return f"{d.day} de {mes} de {d.year}"


def somente_digitos(s: str) -> str:
    return re.sub(r"\D", "", s or "")


def formatar_cpf_cnpj(raw: str) -> str:
    d = somente_digitos(raw)
    if len(d) == 11:  # CPF
        return f"{d[0:3]}.{d[3:6]}.{d[6:9]}-{d[9:11]}"
    if len(d) == 14:  # CNPJ
        return f"{d[0:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:14]}"
    return raw.strip()


def numero_por_extenso(n: int) -> str:
    if n < 0:
        return "menos " + numero_por_extenso(-n)
    if n < 10:
        return UNIDADES[n]
    if n < 20:
        return DEZ_A_DEZENOVE[n - 10]
    if n < 100:
        dez = n // 10
        uni = n % 10
        if uni == 0:
            return DEZENAS[dez]
        return f"{DEZENAS[dez]} e {UNIDADES[uni]}"
    if n == 100:
        return "cem"
    if n < 1000:
        cen = n // 100
        rest = n % 100
        if rest == 0:
            return "cem" if cen == 1 else CENTENAS[cen]
        return f"{CENTENAS[cen]} e {numero_por_extenso(rest)}"
    if n < 1_000_000:
        mil = n // 1000
        rest = n % 1000
        mil_txt = "um mil" if mil == 1 else f"{numero_por_extenso(mil)} mil"
        if rest == 0:
            return mil_txt
        conj = " e " if rest < 100 else " "
        return f"{mil_txt}{conj}{numero_por_extenso(rest)}"
    if n < 1_000_000_000:
        milhao = n // 1_000_000
        rest = n % 1_000_000
        milhao_txt = "um milhão" if milhao == 1 else f"{numero_por_extenso(milhao)} milhões"
        if rest == 0:
            return milhao_txt
        conj = " e " if rest < 100 else " "
        return f"{milhao_txt}{conj}{numero_por_extenso(rest)}"
    bilhao = n // 1_000_000_000
    rest = n % 1_000_000_000
    bilhao_txt = "um bilhão" if bilhao == 1 else f"{numero_por_extenso(bilhao)} bilhões"
    if rest == 0:
        return bilhao_txt
    conj = " e " if rest < 100 else " "
    return f"{bilhao_txt}{conj}{numero_por_extenso(rest)}"


def valor_por_extenso_reais(d: Decimal) -> str:
    d = d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    if d < 0:
        return "menos " + valor_por_extenso_reais(-d)
    reais = int(d // 1)
    centavos = int((d - Decimal(reais)) * 100)

    if reais == 0:
        reais_txt = "zero real" if centavos == 0 else "zero reais"
    elif reais == 1:
        reais_txt = "um real"
    else:
        reais_txt = f"{numero_por_extenso(reais)} reais"

    if centavos == 0:
        return reais_txt

    if centavos == 1:
        cent_txt = "um centavo"
    else:
        cent_txt = f"{numero_por_extenso(centavos)} centavos"

    return f"{reais_txt} e {cent_txt}"


# -----------------------------
# Helpers: DOCX replace (python-docx)
# -----------------------------
def _replace_in_paragraph(paragraph, replacements: dict):
    """
    replacements: {placeholder_string: value_string}
    Faz replace mesmo que o placeholder esteja quebrado em vários runs.
    """
    if not paragraph.runs:
        return

    for placeholder, value in replacements.items():
        while True:
            full_text = "".join(r.text for r in paragraph.runs)
            if placeholder not in full_text:
                break

            start = full_text.find(placeholder)
            end = start + len(placeholder)

            # map runs
            run_starts = []
            pos = 0
            for r in paragraph.runs:
                run_starts.append(pos)
                pos += len(r.text)

            start_run = None
            end_run = None
            for i, start_pos in enumerate(run_starts):
                run_end = start_pos + len(paragraph.runs[i].text)
                if start_run is None and start < run_end:
                    start_run = i
                if start_run is not None and end <= run_end:
                    end_run = i
                    break

            if start_run is None:
                break
            if end_run is None:
                end_run = len(paragraph.runs) - 1

            start_pos = run_starts[start_run]
            # prefix in first run
            prefix = paragraph.runs[start_run].text[: max(0, start - start_pos)]
            # suffix in last run
            suffix = paragraph.runs[end_run].text[max(0, end - run_starts[end_run]):]

            paragraph.runs[start_run].text = prefix + value + suffix
            for j in range(start_run + 1, end_run + 1):
                paragraph.runs[j].text = ""


def _iter_all_paragraphs(doc: Document):
    # body
    for p in doc.paragraphs:
        yield p
    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # headers/footers
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            yield p
        for t in sec.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for p in sec.footer.paragraphs:
            yield p
        for t in sec.footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def replace_text_in_doc(doc: Document, context: dict):
    """
    Substitui {{ KEY }} e variantes.
    """
    replacements = {}
    for k, v in context.items():
        if v is None:
            v = ""
        if not isinstance(v, str):
            v = str(v)
        replacements[f"{{{{ {k} }}}}"] = v
        replacements[f"{{{{{k}}}}}"] = v
        replacements[f"{{{{ {k}}}}}"] = v
        replacements[f"{{{{{k} }}}}"] = v

    for p in _iter_all_paragraphs(doc):
        _replace_in_paragraph(p, replacements)


def replace_image_placeholder(doc: Document, key: str, image_path: str, max_w_mm: float = 120, max_h_mm: float = 45):
    """
    Procura {{ IMAGEM }} e insere a imagem mantendo proporção,
    encaixando em max_w_mm x max_h_mm.
    """
    placeholder_variants = [
        f"{{{{ {key} }}}}",
        f"{{{{{key}}}}}",
        f"{{{{ {key}}}}}",
        f"{{{{{key} }}}}",
    ]

    img = Image.open(image_path)
    w_px, h_px = img.size
    scale = min(max_w_mm / w_px, max_h_mm / h_px)
    w_mm = w_px * scale
    h_mm = h_px * scale

    for p in _iter_all_paragraphs(doc):
        full_text = "".join(r.text for r in p.runs)
        found = None
        for ph in placeholder_variants:
            if ph in full_text:
                found = ph
                break
        if not found:
            continue

        # remove placeholder
        _replace_in_paragraph(p, {found: ""})
        # add picture
        run = p.add_run()
        run.add_picture(image_path, width=Mm(w_mm), height=Mm(h_mm))
        return  # só 1 vez


# -----------------------------
# Helpers: PDF conversion
# -----------------------------
def docx_to_pdf_bytes(docx_path: str) -> bytes:
    """
    Converte DOCX -> PDF usando LibreOffice (soffice).
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        proc = subprocess.run(
            [
                "soffice",
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                docx_path,
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        if proc.returncode != 0:
            raise RuntimeError(f"Falha ao converter para PDF:\n{proc.stdout}")

        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(tmpdir, base + ".pdf")
        if not os.path.exists(pdf_path):
            # fallback
            pdfs = [p for p in os.listdir(tmpdir) if p.lower().endswith(".pdf")]
            if not pdfs:
                raise RuntimeError("PDF não foi gerado pelo LibreOffice.")
            pdf_path = os.path.join(tmpdir, pdfs[0])

        with open(pdf_path, "rb") as f:
            return f.read()


# -----------------------------
# Database
# -----------------------------
def get_database_url() -> str | None:
    url = os.environ.get("DATABASE_URL")
    if not url:
        return None
    # psycopg2 às vezes reclama com postgres://
    if url.startswith("postgres://"):
        url = "postgresql://" + url[len("postgres://"):]
    return url


def db_conn():
    url = get_database_url()
    if not url:
        raise RuntimeError("DATABASE_URL não encontrada (adicione o PostgreSQL no Railway).")
    conn = psycopg2.connect(url)
    conn.autocommit = True
    return conn


def ensure_schema():
    with db_conn() as conn, conn.cursor() as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS propostas (
            id SERIAL PRIMARY KEY,
            cliente TEXT NOT NULL,
            cpf TEXT NOT NULL,
            modelo TEXT NOT NULL,
            franquia INTEGER NOT NULL,
            valor NUMERIC(12,2) NOT NULL,
            pdf BYTEA NOT NULL,
            created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        """)
        # garantir colunas (caso tabela antiga)
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS created_at TIMESTAMPTZ NOT NULL DEFAULT NOW();")
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS pdf BYTEA;")
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS valor NUMERIC(12,2);")
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS franquia INTEGER;")
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS cliente TEXT;")
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS cpf TEXT;")
        cur.execute("ALTER TABLE propostas ADD COLUMN IF NOT EXISTS modelo TEXT;")


def limpar_propostas_expiradas():
    with db_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "DELETE FROM propostas WHERE created_at < NOW() - INTERVAL %s;",
            (f"{RETENCAO_DIAS} days",),
        )


def salvar_proposta(cliente: str, cpf: str, modelo: str, franquia: int, valor: Decimal, pdf_bytes: bytes) -> int:
    with db_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO propostas (cliente, cpf, modelo, franquia, valor, pdf, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, NOW())
            RETURNING id;
            """,
            (cliente, cpf, modelo, franquia, valor, psycopg2.Binary(pdf_bytes)),
        )
        return int(cur.fetchone()[0])


def listar_propostas(limit: int = 50):
    with db_conn() as conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT id, cliente, cpf, modelo, franquia, valor, created_at
            FROM propostas
            ORDER BY created_at DESC
            LIMIT %s;
            """,
            (limit,),
        )
        return cur.fetchall()


def buscar_proposta_pdf(pid: int):
    with db_conn() as conn, conn.cursor() as cur:
        cur.execute("SELECT cliente, pdf FROM propostas WHERE id = %s;", (pid,))
        row = cur.fetchone()
        if not row:
            return None
        cliente, pdf = row
        # pdf pode vir como memoryview
        if isinstance(pdf, memoryview):
            pdf = pdf.tobytes()
        return {"cliente": cliente, "pdf": pdf}


def buscar_proposta_dados(pid: int):
    with db_conn() as conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            "SELECT id, cliente, cpf, modelo, franquia, valor FROM propostas WHERE id = %s;",
            (pid,),
        )
        return cur.fetchone()


def deletar_proposta(pid: int):
    with db_conn() as conn, conn.cursor() as cur:
        cur.execute("DELETE FROM propostas WHERE id = %s;", (pid,))


# -----------------------------
# Routes
# -----------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/proposta", methods=["GET", "POST"])
def proposta():
    if request.method == "GET":
        return render_template("proposta.html", error=None)

    try:
        cliente = (request.form.get("cliente") or "").strip()
        cpf_raw = (request.form.get("cpf") or "").strip()
        modelo = (request.form.get("modelo") or "").strip()
        franquia = parse_int(request.form.get("franquia") or "")
        valor = parse_valor_decimal(request.form.get("valor") or "")

        if not cliente or not cpf_raw or not modelo:
            raise ValueError("Preencha cliente, CPF/CNPJ e modelo.")

        cpf_digits = somente_digitos(cpf_raw)

        imagem_file = request.files.get("imagem")
        if not imagem_file or not imagem_file.filename:
            raise ValueError("Anexe a imagem do equipamento.")

        filename = secure_filename(imagem_file.filename)
        ext = os.path.splitext(filename)[1].lower() or ".png"

        with tempfile.TemporaryDirectory() as tmpdir:
            img_path = os.path.join(tmpdir, "equipamento" + ext)
            imagem_file.save(img_path)

            doc = Document(PROPOSTA_TEMPLATE)
            hoje = datetime.now(APP_TZ).date()

            ctx = {
                "CLIENTE": cliente,
                "CPF": formatar_cpf_cnpj(cpf_digits),
                "MODELO": modelo,
                "FRANQUIA": str(franquia),
                "VALOR": f"R$ {formatar_decimal_ptbr(valor)} ({capitalizar_primeira(valor_por_extenso_reais(valor))})",
                "DATA": data_por_extenso(hoje, mes_capitalizado=True),
            }

            replace_text_in_doc(doc, ctx)
            replace_image_placeholder(doc, "IMAGEM", img_path, max_w_mm=120, max_h_mm=45)

            docx_out = os.path.join(tmpdir, "proposta_gerada.docx")
            doc.save(docx_out)
            pdf_bytes = docx_to_pdf_bytes(docx_out)

        try:
            ensure_schema()
            salvar_proposta(cliente, cpf_digits, modelo, franquia, valor, pdf_bytes)
        except Exception:
            pass

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"Proposta ({cliente}).pdf",
        )

    except Exception as e:
        return render_template("proposta.html", error=str(e))


@app.route("/contrato", methods=["GET", "POST"])
def contrato():
    prefill = {}
    proposta_id = request.args.get("proposta_id")
    if proposta_id:
        try:
            pid = int(proposta_id)
            row = buscar_proposta_dados(pid)
            if row:
                prefill = {
                    "denominacao": row["cliente"],
                    "cpf_cnpj": row["cpf"],
                    "equipamento": row["modelo"],
                    "franquia": str(row["franquia"]),
                    "valor_mensal": str(row["valor"]),
                }
        except Exception:
            prefill = {}

    if request.method == "GET":
        return render_template("contrato.html", prefill=prefill, error=None)

    try:
        denominacao = (request.form.get("denominacao") or "").strip()
        cpf_cnpj_raw = (request.form.get("cpf_cnpj") or "").strip()
        endereco = (request.form.get("endereco") or "").strip()
        telefone = (request.form.get("telefone") or "").strip()
        email = (request.form.get("email") or "").strip()
        equipamento = (request.form.get("equipamento") or "").strip()
        acessorios = (request.form.get("acessorios") or "").strip()
        data_inicio_raw = (request.form.get("data_inicio") or "").strip()
        data_termino_raw = (request.form.get("data_termino") or "").strip()
        franquia = parse_int(request.form.get("franquia") or "")
        valor_mensal = parse_valor_decimal(request.form.get("valor_mensal") or "")

        if not all([denominacao, cpf_cnpj_raw, endereco, telefone, email, equipamento, acessorios, data_inicio_raw, data_termino_raw]):
            raise ValueError("Preencha todos os campos do contrato.")

        di = parse_data_digitos(data_inicio_raw)
        dt = parse_data_digitos(data_termino_raw)
        hoje = datetime.now(APP_TZ).date()

        cpf_digits = somente_digitos(cpf_cnpj_raw)

        doc = Document(CONTRATO_TEMPLATE)
        ctx = {
            "DENOMINACAO": denominacao,
            "CPF_CNPJ": " " + formatar_cpf_cnpj(cpf_digits),
            "ENDERECO": " " + endereco,
            "TELEFONE": " " + telefone,
            "EMAIL": " " + email,
            "EQUIPAMENTO": equipamento,
            "ACESSORIOS": acessorios,
            "DATA_INICIO": data_por_extenso(di, mes_capitalizado=False),
            "DATA_TERMINO": data_por_extenso(dt, mes_capitalizado=False),
            "FRANQUIA_FORMATADA": formatar_inteiro_ptbr(franquia),
            "FRANQUIA_EXTENSO": numero_por_extenso(franquia),
            "VALOR_MENSAL_FORMATADO": formatar_decimal_ptbr(valor_mensal),
            "VALOR_MENSAL_EXTENSO": valor_por_extenso_reais(valor_mensal),
            "DATA_ASSINATURA": data_por_extenso(hoje, mes_capitalizado=False),
        }
        replace_text_in_doc(doc, ctx)

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_out = os.path.join(tmpdir, "contrato_gerado.docx")
            doc.save(docx_out)
            pdf_bytes = docx_to_pdf_bytes(docx_out)

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"Contrato ({denominacao}).pdf",
        )

    except Exception as e:
        prefill_post = dict(request.form)
        return render_template("contrato.html", prefill=prefill_post, error=str(e))


@app.route("/recentes")
def recentes():
    try:
        ensure_schema()
        limpar_propostas_expiradas()
        propostas = listar_propostas()
        return render_template("recentes.html", propostas=propostas, error=None)
    except Exception as e:
        return render_template("recentes.html", propostas=[], error=str(e))


@app.route("/propostas/<int:pid>/pdf")
def proposta_pdf(pid: int):
    row = buscar_proposta_pdf(pid)
    if not row:
        return "Proposta não encontrada.", 404

    download = request.args.get("download") == "1"
    cliente = row["cliente"] or "cliente"
    pdf_bytes = row["pdf"]

    return send_file(
        io.BytesIO(pdf_bytes),
        mimetype="application/pdf",
        as_attachment=download,
        download_name=f"Proposta ({cliente}).pdf",
    )


@app.route("/propostas/<int:pid>/delete", methods=["POST"])
def proposta_delete(pid: int):
    try:
        deletar_proposta(pid)
    except Exception:
        pass
    return redirect(url_for("recentes"))


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "8000"))
    app.run(host="0.0.0.0", port=port)
